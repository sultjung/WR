#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document

REQUIRED_HEADINGS = [
    "1. 이라크 국내 상황",
    "1) 정국 / 치안",
    "· 정치권 동향",
    "· 이라크 주간 테러 상황",
    "2) 경제",
    "· 국제유가 관련 동향",
    "2. 국제사회",
    "3. 그룹 / 건설에 미치는 영향",
]
SECTION_NAMES = ("politicsItems", "securityItems", "economyItems", "internationalItems")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", default="work/weekly-report.docx")
    parser.add_argument("--content", default="work/report-content.json")
    parser.add_argument("--metadata", default="work/report-metadata.json")
    return parser.parse_args()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def validate_report_article_structure(content: dict) -> dict:
    request = content["request"]
    report = content["report"]
    selected = set(request["selectedArticleIds"])
    article_index = {item["articleId"]: item for item in content.get("selectedArticleIndex", [])}
    if set(article_index) != selected:
        raise RuntimeError("selectedArticleIndex does not match selected article IDs")

    usage: Counter[str] = Counter()
    assignment: dict[str, str] = {}
    section_dates: dict[str, list[str]] = {}
    for section in SECTION_NAMES:
        dates: list[str] = []
        for item_index, item in enumerate(report.get(section, [])):
            article_ids = item.get("articleIds", [])
            if len(article_ids) != len(set(article_ids)):
                raise RuntimeError(f"duplicate article ID inside one item: {item.get('headline', '')}")
            item_dates: list[str] = []
            for article_id in article_ids:
                if article_id not in selected:
                    raise RuntimeError(f"unselected article used in report: {article_id}")
                meta = article_index[article_id]
                if meta.get("targetSection") != section:
                    raise RuntimeError(f"article {article_id} belongs to {meta.get('targetSection')} but appears in {section}")
                usage[article_id] += 1
                assignment[article_id] = f"{section}:{item_index}"
                item_dates.append(str(meta.get("publishedDate", "")))
            if item_dates:
                dates.append(min(item_dates))
        if dates != sorted(dates):
            raise RuntimeError(f"items are not chronological inside {section}: {dates}")
        section_dates[section] = dates

    missing = sorted(article_id for article_id in selected if usage[article_id] == 0)
    duplicated = sorted(article_id for article_id in selected if usage[article_id] > 1)
    if missing or duplicated:
        raise RuntimeError(f"selected article usage mismatch; missing={missing}, duplicated={duplicated}")

    for cluster in content.get("reportClusters", []):
        article_ids = cluster.get("articleIds", [])
        if len(article_ids) < 2:
            continue
        locations = {assignment.get(article_id) for article_id in article_ids}
        locations.discard(None)
        if len(locations) != 1:
            raise RuntimeError(f"similar-article cluster {cluster.get('clusterId')} was split across report items: {sorted(locations)}")

    return {
        "selected": selected,
        "sectionDates": section_dates,
        "clusterCount": len(content.get("reportClusters", [])),
        "mergedClusterCount": sum(1 for cluster in content.get("reportClusters", []) if len(cluster.get("articleIds", [])) > 1),
    }


def validate(docx_path: Path, content_path: Path) -> dict:
    if not docx_path.exists():
        raise RuntimeError("Word output does not exist")
    if docx_path.stat().st_size < 25_000:
        raise RuntimeError(f"Word output is unexpectedly small: {docx_path.stat().st_size} bytes")

    content = json.loads(content_path.read_text(encoding="utf-8"))
    structure = validate_report_article_structure(content)
    document = Document(docx_path)
    text = all_text(document)

    if "{{" in text or "}}" in text:
        raise RuntimeError("unresolved template placeholder remains in Word output")
    if re.search(r"\bAI\b|프롬프트|자동\s*생성|선택\s*기사|모델", text, flags=re.I):
        raise RuntimeError("production-process wording leaked into the report")
    for heading in REQUIRED_HEADINGS:
        if heading not in text:
            raise RuntimeError(f"required heading missing: {heading}")

    request = content["request"]
    report = content["report"]
    oil = content["oil"]
    expected_title = "건설, 이라크 주간 종합 상황보고"
    if expected_title not in text:
        raise RuntimeError("report title is missing")
    report_date = request["reportDate"]
    y, m, d = (int(part) for part in report_date.split("-"))
    if f"{y}. {m}. {d}." not in text:
        raise RuntimeError("report date is missing or malformed")

    if len(document.tables) < 2:
        raise RuntimeError(f"expected at least 2 tables, found {len(document.tables)}")
    terror = report["terrorStats"]
    terror_sum = sum(terror[key] for key in ("armedAttack", "ied", "assassination", "protest", "shooting", "suicideBombing"))
    if terror["total"] != terror_sum:
        raise RuntimeError("terror total does not equal category sum")

    for row in oil["prices"]:
        md = f"{int(row['date'][5:7])}.{int(row['date'][8:10])}"
        expected_values = [md, f"${row['dubai']:.2f}", f"${row['brent']:.2f}", f"${row['wti']:.2f}"]
        for value in expected_values:
            if value not in text:
                raise RuntimeError(f"oil table value missing: {value}")

    metadata = {
        "pipelineVersion": content.get("pipelineVersion"),
        "generatedAt": content.get("generatedAt"),
        "reportDate": report_date,
        "periodStart": request["periodStart"],
        "periodEnd": request["periodEnd"],
        "issueNumber": request.get("issueNumber"),
        "selectedArticleCount": len(structure["selected"]),
        "clusterCount": structure["clusterCount"],
        "mergedClusterCount": structure["mergedClusterCount"],
        "sectionDates": structure["sectionDates"],
        "model": content.get("model"),
        "responseId": content.get("responseId"),
        "oilSource": oil.get("source"),
        "oilSourceUrl": oil.get("sourceUrl"),
        "oilTargetDates": oil.get("targetDates"),
        "docxSha256": sha256(docx_path),
        "docxBytes": docx_path.stat().st_size,
        "tableCount": len(document.tables),
        "paragraphCount": len(document.paragraphs),
    }
    return metadata


def main() -> int:
    args = parse_args()
    docx_path = Path(args.docx)
    content_path = Path(args.content)
    metadata_path = Path(args.metadata)
    metadata = validate(docx_path, content_path)
    metadata_path.parent.mkdir(parents=True, exist_ok=True)
    metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"[report-validate] passed bytes={metadata['docxBytes']} tables={metadata['tableCount']} paragraphs={metadata['paragraphCount']} clusters={metadata['clusterCount']} merged={metadata['mergedClusterCount']} sha256={metadata['docxSha256'][:12]}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as error:  # noqa: BLE001
        print(f"[report-validate] ERROR: {error}", file=sys.stderr)
        raise
