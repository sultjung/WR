#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

FONT_NAME = "Batang"


def set_east_asia_font(run, font_name=FONT_NAME):
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, size, bold=False):
    style.font.name = FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_east_asia_font(run)
    run.font.size = Pt(10)


def create_template(output: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    set_style_font(normal, 11.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.15

    style_specs = {
        "Report Title": (15, True, 0, 8),
        "Report Date": (11.5, False, 0, 12),
        "Report Heading 1": (14, True, 7, 4),
        "Report Heading 2": (12.5, True, 6, 3),
        "Report Category": (12, True, 4, 2),
        "Report Item": (11.5, False, 2, 1),
        "Report Detail": (11, False, 0, 1),
        "Report Implication": (11, False, 0, 2),
        "Report Impact": (11.5, False, 2, 2),
    }
    for name, (size, bold, before, after) in style_specs.items():
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size, bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.widow_control = True

    doc.styles["Report Item"].paragraph_format.left_indent = Mm(4)
    doc.styles["Report Detail"].paragraph_format.left_indent = Mm(9)
    doc.styles["Report Implication"].paragraph_format.left_indent = Mm(9)
    doc.styles["Report Impact"].paragraph_format.left_indent = Mm(4)

    title = doc.add_paragraph(style="Report Title")
    title_run = title.add_run("건설, 이라크 주간 종합 상황보고({{PERIOD}})")
    set_east_asia_font(title_run)
    title_run.underline = True

    report_date = doc.add_paragraph(style="Report Date")
    report_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = report_date.add_run("{{REPORT_DATE}}")
    set_east_asia_font(date_run)

    sample = doc.add_paragraph(style="Report Heading 1")
    run = sample.add_run("1. 이라크 국내 상황")
    set_east_asia_font(run)
    placeholder = doc.add_paragraph(style="Report Item")
    run = placeholder.add_run("{{REPORT_BODY}}")
    set_east_asia_font(run)

    add_page_field(section.footer.paragraphs[0])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", default="templates/weekly-report-template.docx")
    args = parser.parse_args()
    create_template(Path(args.output))
    print(f"created {args.output}")


if __name__ == "__main__":
    main()
