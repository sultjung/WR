#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

FONT_NAME = "Batang"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", default="templates/weekly-report-template.docx")
    parser.add_argument("--content", default="work/report-content.json")
    parser.add_argument("--output", default="work/weekly-report.docx")
    return parser.parse_args()


def set_font(run, size: float | None = None, bold: bool | None = None, underline: bool | None = None):
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.underline = underline


def clear_body_keep_section(document: Document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def format_period(start: str, end: str) -> str:
    sy, sm, sd = (int(part) for part in start.split("-"))
    ey, em, ed = (int(part) for part in end.split("-"))
    return f"΄{str(sy)[-2:]}.{sm}.{sd} ~ ΄{str(ey)[-2:]}.{em}.{ed}"


def format_report_date(value: str) -> str:
    year, month, day = (int(part) for part in value.split("-"))
    return f"{year}. {month}. {day}."


def format_md(value: str) -> str:
    _, month, day = (int(part) for part in value.split("-"))
    return f"{month}.{day}"


def add_paragraph(document: Document, text: str, style: str, *, keep_with_next: bool = False, align=None):
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_font(run)
    return paragraph


def set_cell_text(cell, text: str, *, bold=False, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    set_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, width in (("top", 70), ("left", 70), ("bottom", 70), ("right", 70)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill="E7E6E6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for index, value in enumerate(headers):
        set_cell_text(header.cells[index], value, bold=True, size=9.2)
        shade_cell(header.cells[index])
    for values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 35 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(row.cells[index], value, size=9.2, align=align)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Mm(width)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)
    return table


def add_report_item(document: Document, item: dict):
    date_label = str(item.get("dateLabel", "")).strip()
    headline = str(item.get("headline", "")).strip()
    main = f"- {date_label}, {headline}" if date_label else f"- {headline}"
    add_paragraph(document, main, "Report Item", keep_with_next=bool(item.get("details") or item.get("tableHeaders") or item.get("implication")))
    headers = [str(value).strip() for value in item.get("tableHeaders", []) if str(value).strip()]
    rows = [[str(cell).strip() for cell in row] for row in item.get("tableRows", [])]
    if headers and rows:
        column_count = len(headers)
        if column_count == 3:
            widths = [18, 42, 110]
        elif column_count == 2:
            widths = [45, 125]
        elif column_count == 4:
            widths = [20, 45, 50, 55]
        else:
            widths = [170 / column_count] * column_count
        add_table(document, headers, rows, widths)
    for detail in item.get("details", []):
        detail_text = str(detail).strip()
        if detail_text:
            add_paragraph(document, f"* {detail_text}", "Report Detail")
    implication = str(item.get("implication", "")).strip()
    if implication:
        add_paragraph(document, f"☞ {implication}", "Report Implication")


def add_items(document: Document, items: list[dict]):
    for item in items:
        add_report_item(document, item)


def build_document(template_path: Path, content: dict, output_path: Path):
    document = Document(template_path)
    clear_body_keep_section(document)
    request = content["request"]
    report = content["report"]
    oil = content["oil"]

    title = add_paragraph(
        document,
        f"건설, 이라크 주간 종합 상황보고({format_period(request['periodStart'], request['periodEnd'])})",
        "Report Title",
    )
    for run in title.runs:
        set_font(run, size=15, bold=True, underline=True)

    report_date = add_paragraph(document, format_report_date(request["reportDate"]), "Report Date", align=WD_ALIGN_PARAGRAPH.RIGHT)
    for run in report_date.runs:
        set_font(run, size=11.5)

    add_paragraph(document, "1. 이라크 국내 상황", "Report Heading 1", keep_with_next=True)
    add_paragraph(document, "1) 정국 / 치안", "Report Heading 2", keep_with_next=True)
    add_paragraph(document, "· 정치권 동향", "Report Category", keep_with_next=True)
    add_items(document, report.get("politicsItems", []))

    add_paragraph(document, "· 이라크 주간 테러 상황", "Report Category", keep_with_next=True)
    stats = report["terrorStats"]
    add_table(
        document,
        ["구분", "계", "무장세력공격", "IED", "암 살", "시 위", "총 격", "자살폭탄테러"],
        [["건수", stats["total"], stats["armedAttack"], stats["ied"], stats["assassination"], stats["protest"], stats["shooting"], stats["suicideBombing"]]],
        [18, 16, 29, 17, 20, 20, 20, 30],
    )
    add_items(document, report.get("securityItems", []))

    add_paragraph(document, "2) 경제", "Report Heading 2", keep_with_next=True)
    add_paragraph(document, "· 국제유가 관련 동향", "Report Category", keep_with_next=True)
    add_items(document, report.get("economyItems", []))
    oil_rows = [
        [format_md(row["date"]), f"${row['dubai']:.2f}", f"${row['brent']:.2f}", f"${row['wti']:.2f}"]
        for row in oil["prices"]
    ]
    add_table(document, ["구 분", "두바이유", "브렌트유", "서부텍사스유(WTI)"], oil_rows, [34, 42, 42, 52])

    add_paragraph(document, "2. 국제사회", "Report Heading 1", keep_with_next=True)
    topic = str(report.get("internationalTopic", "")).strip() or "주요 국제정세"
    add_paragraph(document, f"· {topic}", "Report Category", keep_with_next=True)
    add_items(document, report.get("internationalItems", []))

    add_paragraph(document, "3. 그룹 / 건설에 미치는 영향", "Report Heading 1", keep_with_next=True)
    for impact in report.get("impactItems", []):
        text = str(impact).strip()
        if text:
            add_paragraph(document, f"· {text}", "Report Impact")

    # Prevent stray placeholder text or accidental blank leading pages.
    for paragraph in document.paragraphs:
        if re.search(r"\{\{.+?\}\}", paragraph.text):
            raise RuntimeError(f"unresolved template placeholder: {paragraph.text}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    template_path = Path(args.template)
    content_path = Path(args.content)
    output_path = Path(args.output)
    content = json.loads(content_path.read_text(encoding="utf-8"))
    build_document(template_path, content, output_path)
    print(f"[report-docx] created {output_path} ({output_path.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
